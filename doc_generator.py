import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

def generate_cv_docx(cv_data: dict, output_path: str):
    """Generates a professional tailored CV in DOCX format."""
    doc = docx.Document()
    
    # Margin settings
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    
    # Title / Header
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(2)
    
    title_run = title_p.add_run("HUZAIFA NAEEM\n")
    title_run.font.name = "Arial"
    title_run.font.size = Pt(18)
    title_run.bold = True
    
    subtitle_run = title_p.add_run("Senior Product Designer\n")
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(12)
    subtitle_run.bold = True
    
    contact_run = title_p.add_run("+92 324 2237858 | huzaifnaeem06@gmail.com | LinkedIn | Behance | Karachi, Pakistan | Remote")
    contact_run.font.name = "Arial"
    contact_run.font.size = Pt(9)
    
    def add_heading(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        run = h.add_run(text.upper())
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(11)
        
    # Profile
    add_heading("Profile")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p_run = p.add_run(cv_data.get("profile", ""))
    p_run.font.name = "Arial"
    p_run.font.size = Pt(10)
    
    # Experience
    add_heading("Experience")
    for exp in cv_data.get("experience", []):
        p_title = doc.add_paragraph()
        p_title.paragraph_format.space_before = Pt(6)
        p_title.paragraph_format.space_after = Pt(2)
        r_title = p_title.add_run(f"{exp.get('role')}  —  {exp.get('company')}\n")
        r_title.bold = True
        r_title.font.name = "Arial"
        r_title.font.size = Pt(10)
        
        r_meta = p_title.add_run(f"{exp.get('duration')} | {exp.get('location')} | {exp.get('type')}")
        r_meta.italic = True
        r_meta.font.name = "Arial"
        r_meta.font.size = Pt(9)
        
        for bullet in exp.get("bullets", []):
            bp = doc.add_paragraph(style='List Bullet')
            bp.paragraph_format.space_after = Pt(2)
            brun = bp.add_run(bullet)
            brun.font.name = "Arial"
            brun.font.size = Pt(9.5)
            
    # Key Projects
    add_heading("Key Projects")
    for proj in cv_data.get("projects", []):
        pp = doc.add_paragraph()
        pp.paragraph_format.space_before = Pt(4)
        pp.paragraph_format.space_after = Pt(2)
        prun = pp.add_run(proj.get('name', '') + "\n")
        prun.bold = True
        prun.font.name = "Arial"
        prun.font.size = Pt(10)
        
        pdesc = pp.add_run(proj.get('details', ''))
        pdesc.font.name = "Arial"
        pdesc.font.size = Pt(9.5)
        
    # Skills, Tools, Education, Languages
    add_heading("Core Skills")
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)
    srun = sp.add_run(", ".join(cv_data.get("skills", [])))
    srun.font.name = "Arial"
    srun.font.size = Pt(9.5)
    
    add_heading("Tools & Software")
    tp = doc.add_paragraph()
    tp.paragraph_format.space_after = Pt(4)
    trun = tp.add_run(", ".join(cv_data.get("tools", [])))
    trun.font.name = "Arial"
    trun.font.size = Pt(9.5)
    
    add_heading("Education")
    edup = doc.add_paragraph()
    edup.paragraph_format.space_after = Pt(2)
    edurun = edup.add_run("B.S. Computer Science (In Progress)  —  Muhammad Ali Jinnah University, Karachi\n")
    edurun.bold = True
    edurun.font.name = "Arial"
    edurun.font.size = Pt(10)
    edumeta = edup.add_run("Feb 2026 – Feb 2030")
    edumeta.italic = True
    edumeta.font.name = "Arial"
    edumeta.font.size = Pt(9)
    
    add_heading("Languages")
    langp = doc.add_paragraph()
    langrun = langp.add_run("Urdu (Native) | English (Professional)")
    langrun.font.name = "Arial"
    langrun.font.size = Pt(9.5)
    
    doc.save(output_path)

def generate_cover_letter_docx(cl_data: dict, output_path: str):
    """Generates a professional cover letter in DOCX format."""
    doc = docx.Document()
    
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    
    # Header
    header_p = doc.add_paragraph()
    hrun1 = header_p.add_run("HUZAIFA NAEEM\n")
    hrun1.bold = True
    hrun1.font.size = Pt(14)
    hrun1.font.name = "Arial"
    
    hrun2 = header_p.add_run("Senior Product Designer\n+92 324 2237858 | huzaifnaeem06@gmail.com\n\n")
    hrun2.font.size = Pt(10)
    hrun2.font.name = "Arial"
    
    # Recipient & Date
    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_after = Pt(12)
    mrun = p_meta.add_run(f"Date: {cl_data.get('date', 'July 16, 2026')}\nTo: Hiring Team\nCompany: {cl_data.get('company', '')}\n")
    mrun.font.size = Pt(10)
    mrun.font.name = "Arial"
    
    # Body
    body_text = cl_data.get("letter_body", "")
    for paragraph in body_text.split('\n\n'):
        if paragraph.strip():
            bp = doc.add_paragraph()
            bp.paragraph_format.space_after = Pt(8)
            brun = bp.add_run(paragraph.strip())
            brun.font.size = Pt(10.5)
            brun.font.name = "Arial"
        
    doc.save(output_path)

def generate_cv_pdf(cv_data: dict, output_path: str):
    """Generates a professional tailored CV in PDF format."""
    doc = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        leftMargin=54,
        rightMargin=54,
        topMargin=54,
        bottomMargin=54
    )
    
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'HeaderTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=18,
        leading=22,
        alignment=1,
        spaceAfter=2
    )
    
    subtitle_style = ParagraphStyle(
        'HeaderSubtitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=12,
        leading=14,
        alignment=1,
        spaceAfter=4
    )
    
    contact_style = ParagraphStyle(
        'HeaderContact',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        leading=11,
        alignment=1,
        spaceAfter=10
    )
    
    section_heading_style = ParagraphStyle(
        'SectionHeading',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=11,
        leading=14,
        textColor=colors.HexColor('#222222'),
        spaceBefore=10,
        spaceAfter=4,
        keepWithNext=True
    )
    
    body_style = ParagraphStyle(
        'CVBody',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13,
        spaceAfter=6
    )
    
    bullet_style = ParagraphStyle(
        'CVBullet',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        leading=12.5,
        leftIndent=15,
        firstLineIndent=-10,
        spaceAfter=3
    )
    
    exp_title_style = ParagraphStyle(
        'ExpTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=9.5,
        leading=12,
        spaceBefore=4,
        spaceAfter=1,
        keepWithNext=True
    )
    
    exp_meta_style = ParagraphStyle(
        'ExpMeta',
        parent=styles['Normal'],
        fontName='Helvetica-Oblique',
        fontSize=8.5,
        leading=10,
        spaceAfter=3,
        keepWithNext=True
    )

    story = []
    
    # Header
    story.append(Paragraph("HUZAIFA NAEEM", title_style))
    story.append(Paragraph("Senior Product Designer", subtitle_style))
    story.append(Paragraph("+92 324 2237858 | huzaifnaeem06@gmail.com | LinkedIn | Behance | Karachi, Pakistan | Remote", contact_style))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#CCCCCC'), spaceAfter=8))
    
    # Profile
    story.append(Paragraph("PROFILE", section_heading_style))
    story.append(Paragraph(cv_data.get("profile", ""), body_style))
    
    # Experience
    story.append(Paragraph("EXPERIENCE", section_heading_style))
    for exp in cv_data.get("experience", []):
        story.append(Paragraph(f"{exp.get('role', '')} &mdash; {exp.get('company', '')}", exp_title_style))
        story.append(Paragraph(f"{exp.get('duration', '')} | {exp.get('location', '')} | {exp.get('type', '')}", exp_meta_style))
        for bullet in exp.get("bullets", []):
            story.append(Paragraph(f"&bull;&nbsp;&nbsp;{bullet}", bullet_style))
            
    # Key Projects
    story.append(Paragraph("KEY PROJECTS", section_heading_style))
    for proj in cv_data.get("projects", []):
        story.append(Paragraph(proj.get('name', ''), exp_title_style))
        story.append(Paragraph(proj.get('details', ''), body_style))
        
    # Skills
    story.append(Paragraph("CORE SKILLS", section_heading_style))
    story.append(Paragraph(", ".join(cv_data.get("skills", [])), body_style))
    
    # Tools
    story.append(Paragraph("TOOLS &amp; SOFTWARE", section_heading_style))
    story.append(Paragraph(", ".join(cv_data.get("tools", [])), body_style))
    
    # Education
    story.append(Paragraph("EDUCATION", section_heading_style))
    story.append(Paragraph("B.S. Computer Science (In Progress) &mdash; Muhammad Ali Jinnah University, Karachi", exp_title_style))
    story.append(Paragraph("Feb 2026 &ndash; Feb 2030", exp_meta_style))
    
    # Languages
    story.append(Paragraph("LANGUAGES", section_heading_style))
    story.append(Paragraph("Urdu (Native) | English (Professional)", body_style))
    
    doc.build(story)

def generate_cover_letter_pdf(cl_data: dict, output_path: str):
    """Generates a professional cover letter in PDF format."""
    doc = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        leftMargin=72,
        rightMargin=72,
        topMargin=72,
        bottomMargin=72
    )
    
    styles = getSampleStyleSheet()
    
    header_style = ParagraphStyle(
        'CLHeader',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=14,
        leading=17,
        spaceAfter=2
    )
    
    header_sub_style = ParagraphStyle(
        'CLHeaderSub',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=12,
        spaceAfter=15
    )
    
    meta_style = ParagraphStyle(
        'CLMeta',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13,
        spaceAfter=12
    )
    
    body_style = ParagraphStyle(
        'CLBody',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=14.5,
        spaceAfter=10
    )
    
    story = []
    
    story.append(Paragraph("HUZAIFA NAEEM", header_style))
    story.append(Paragraph("Senior Product Designer<br/>+92 324 2237858 | huzaifnaeem06@gmail.com", header_sub_style))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#CCCCCC'), spaceAfter=12))
    
    story.append(Paragraph(f"Date: {cl_data.get('date', 'July 16, 2026')}<br/>To: Hiring Team<br/>Company: {cl_data.get('company', '')}", meta_style))
    
    body_text = cl_data.get("letter_body", "")
    for paragraph in body_text.split('\n\n'):
        if paragraph.strip():
            story.append(Paragraph(paragraph.strip().replace('\n', '<br/>'), body_style))
            
    doc.build(story)

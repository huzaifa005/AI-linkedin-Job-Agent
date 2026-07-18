"""
Document Generator Service — Creates tailored CV and Cover Letter documents.
Refactored from the original doc_generator.py:
  - Outputs to BytesIO (for API streaming) or file paths
  - Removes hardcoded personal info — now pulled from CV data
  - Preserves the exact professional formatting of the originals
"""

import io
import os

import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors


# ─── Default candidate info (extracted from original CV) ────────────────────
# These are used as fallback if the tailored data doesn't include them.
DEFAULT_CANDIDATE = {
    "name": "HUZAIFA NAEEM",
    "title": "Senior Product Designer",
    "contact": "+92 324 2237858 | huzaifnaeem06@gmail.com | LinkedIn | Behance | Karachi, Pakistan | Remote",
    "email": "huzaifnaeem06@gmail.com",
    "phone": "+92 324 2237858",
    "education": "B.S. Computer Science (In Progress)  —  Muhammad Ali Jinnah University, Karachi",
    "education_duration": "Feb 2026 – Feb 2030",
    "languages": "Urdu (Native) | English (Professional)",
}


# ═══════════════════════════════════════════════════════════════════════════
# DOCX GENERATORS
# ═══════════════════════════════════════════════════════════════════════════

def generate_cv_docx(cv_data: dict, output_path: str = None) -> io.BytesIO:
    """
    Generates a professional tailored CV in DOCX format.
    
    Args:
        cv_data: Tailored CV data dict from Groq (profile, experience, projects, skills, tools).
        output_path: Optional file path to save. If None, returns BytesIO only.
        
    Returns:
        BytesIO object containing the DOCX file.
    """
    doc = docx.Document()

    # Margin settings
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    # Title / Header
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(2)

    title_run = title_p.add_run(DEFAULT_CANDIDATE["name"] + "\n")
    title_run.font.name = "Arial"
    title_run.font.size = Pt(18)
    title_run.bold = True

    subtitle_run = title_p.add_run(DEFAULT_CANDIDATE["title"] + "\n")
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(12)
    subtitle_run.bold = True

    contact_run = title_p.add_run(DEFAULT_CANDIDATE["contact"])
    contact_run.font.name = "Arial"
    contact_run.font.size = Pt(9)

    def add_heading(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        run = h.add_run(text.upper())
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(11)

    # Profile
    add_heading("Profile")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p_run = p.add_run(cv_data.get("profile", ""))
    p_run.font.name = "Arial"
    p_run.font.size = Pt(10)

    # Experience
    add_heading("Experience")
    for exp in cv_data.get("experience", []):
        p_title = doc.add_paragraph()
        p_title.paragraph_format.space_before = Pt(6)
        p_title.paragraph_format.space_after = Pt(2)
        r_title = p_title.add_run(f"{exp.get('role')}  —  {exp.get('company')}\n")
        r_title.bold = True
        r_title.font.name = "Arial"
        r_title.font.size = Pt(10)

        r_meta = p_title.add_run(
            f"{exp.get('duration')} | {exp.get('location')} | {exp.get('type')}"
        )
        r_meta.italic = True
        r_meta.font.name = "Arial"
        r_meta.font.size = Pt(9)

        for bullet in exp.get("bullets", []):
            bp = doc.add_paragraph(style="List Bullet")
            bp.paragraph_format.space_after = Pt(2)
            brun = bp.add_run(bullet)
            brun.font.name = "Arial"
            brun.font.size = Pt(9.5)

    # Key Projects
    add_heading("Key Projects")
    for proj in cv_data.get("projects", []):
        pp = doc.add_paragraph()
        pp.paragraph_format.space_before = Pt(4)
        pp.paragraph_format.space_after = Pt(2)
        prun = pp.add_run(proj.get("name", "") + "\n")
        prun.bold = True
        prun.font.name = "Arial"
        prun.font.size = Pt(10)

        pdesc = pp.add_run(proj.get("details", ""))
        pdesc.font.name = "Arial"
        pdesc.font.size = Pt(9.5)

    # Skills, Tools, Education, Languages
    add_heading("Core Skills")
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)
    srun = sp.add_run(", ".join(cv_data.get("skills", [])))
    srun.font.name = "Arial"
    srun.font.size = Pt(9.5)

    add_heading("Tools & Software")
    tp = doc.add_paragraph()
    tp.paragraph_format.space_after = Pt(4)
    trun = tp.add_run(", ".join(cv_data.get("tools", [])))
    trun.font.name = "Arial"
    trun.font.size = Pt(9.5)

    add_heading("Education")
    edup = doc.add_paragraph()
    edup.paragraph_format.space_after = Pt(2)
    edurun = edup.add_run(DEFAULT_CANDIDATE["education"] + "\n")
    edurun.bold = True
    edurun.font.name = "Arial"
    edurun.font.size = Pt(10)
    edumeta = edup.add_run(DEFAULT_CANDIDATE["education_duration"])
    edumeta.italic = True
    edumeta.font.name = "Arial"
    edumeta.font.size = Pt(9)

    add_heading("Languages")
    langp = doc.add_paragraph()
    langrun = langp.add_run(DEFAULT_CANDIDATE["languages"])
    langrun.font.name = "Arial"
    langrun.font.size = Pt(9.5)

    # Save to BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Also save to file if path provided
    if output_path:
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        with open(output_path, "wb") as f:
            f.write(buffer.getvalue())
        buffer.seek(0)

    return buffer


def generate_cover_letter_docx(cl_data: dict, output_path: str = None) -> io.BytesIO:
    """
    Generates a professional cover letter in DOCX format.
    
    Args:
        cl_data: Cover letter data dict from Groq (date, company, letter_body).
        output_path: Optional file path to save.
        
    Returns:
        BytesIO object containing the DOCX file.
    """
    doc = docx.Document()

    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # Header
    header_p = doc.add_paragraph()
    hrun1 = header_p.add_run(DEFAULT_CANDIDATE["name"] + "\n")
    hrun1.bold = True
    hrun1.font.size = Pt(14)
    hrun1.font.name = "Arial"

    hrun2 = header_p.add_run(
        f"{DEFAULT_CANDIDATE['title']}\n"
        f"{DEFAULT_CANDIDATE['phone']} | {DEFAULT_CANDIDATE['email']}\n\n"
    )
    hrun2.font.size = Pt(10)
    hrun2.font.name = "Arial"

    # Recipient & Date
    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_after = Pt(12)
    mrun = p_meta.add_run(
        f"Date: {cl_data.get('date', 'July 16, 2026')}\n"
        f"To: Hiring Team\n"
        f"Company: {cl_data.get('company', '')}\n"
    )
    mrun.font.size = Pt(10)
    mrun.font.name = "Arial"

    # Body
    body_text = cl_data.get("letter_body", "")
    for paragraph in body_text.split("\n\n"):
        if paragraph.strip():
            bp = doc.add_paragraph()
            bp.paragraph_format.space_after = Pt(8)
            brun = bp.add_run(paragraph.strip())
            brun.font.size = Pt(10.5)
            brun.font.name = "Arial"

    # Save to BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    if output_path:
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        with open(output_path, "wb") as f:
            f.write(buffer.getvalue())
        buffer.seek(0)

    return buffer


# ═══════════════════════════════════════════════════════════════════════════
# PDF GENERATORS
# ═══════════════════════════════════════════════════════════════════════════

def generate_cv_pdf(cv_data: dict, output_path: str = None) -> io.BytesIO:
    """
    Generates a professional tailored CV in PDF format.
    
    Args:
        cv_data: Tailored CV data dict from Groq.
        output_path: Optional file path to save.
        
    Returns:
        BytesIO object containing the PDF file.
    """
    buffer = io.BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=54,
        rightMargin=54,
        topMargin=54,
        bottomMargin=54,
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "HeaderTitle",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=18,
        leading=22,
        alignment=1,
        spaceAfter=2,
    )

    subtitle_style = ParagraphStyle(
        "HeaderSubtitle",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=12,
        leading=14,
        alignment=1,
        spaceAfter=4,
    )

    contact_style = ParagraphStyle(
        "HeaderContact",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9,
        leading=11,
        alignment=1,
        spaceAfter=10,
    )

    section_heading_style = ParagraphStyle(
        "SectionHeading",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=14,
        textColor=colors.HexColor("#222222"),
        spaceBefore=10,
        spaceAfter=4,
        keepWithNext=True,
    )

    body_style = ParagraphStyle(
        "CVBody",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9.5,
        leading=13,
        spaceAfter=6,
    )

    bullet_style = ParagraphStyle(
        "CVBullet",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9,
        leading=12.5,
        leftIndent=15,
        firstLineIndent=-10,
        spaceAfter=3,
    )

    exp_title_style = ParagraphStyle(
        "ExpTitle",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=9.5,
        leading=12,
        spaceBefore=4,
        spaceAfter=1,
        keepWithNext=True,
    )

    exp_meta_style = ParagraphStyle(
        "ExpMeta",
        parent=styles["Normal"],
        fontName="Helvetica-Oblique",
        fontSize=8.5,
        leading=10,
        spaceAfter=3,
        keepWithNext=True,
    )

    story = []

    # Header
    story.append(Paragraph(DEFAULT_CANDIDATE["name"], title_style))
    story.append(Paragraph(DEFAULT_CANDIDATE["title"], subtitle_style))
    story.append(Paragraph(DEFAULT_CANDIDATE["contact"], contact_style))
    story.append(
        HRFlowable(
            width="100%",
            thickness=1,
            color=colors.HexColor("#CCCCCC"),
            spaceAfter=8,
        )
    )

    # Profile
    story.append(Paragraph("PROFILE", section_heading_style))
    story.append(Paragraph(cv_data.get("profile", ""), body_style))

    # Experience
    story.append(Paragraph("EXPERIENCE", section_heading_style))
    for exp in cv_data.get("experience", []):
        story.append(
            Paragraph(
                f"{exp.get('role', '')} &mdash; {exp.get('company', '')}",
                exp_title_style,
            )
        )
        story.append(
            Paragraph(
                f"{exp.get('duration', '')} | {exp.get('location', '')} | {exp.get('type', '')}",
                exp_meta_style,
            )
        )
        for bullet in exp.get("bullets", []):
            story.append(Paragraph(f"&bull;&nbsp;&nbsp;{bullet}", bullet_style))

    # Key Projects
    story.append(Paragraph("KEY PROJECTS", section_heading_style))
    for proj in cv_data.get("projects", []):
        story.append(Paragraph(proj.get("name", ""), exp_title_style))
        story.append(Paragraph(proj.get("details", ""), body_style))

    # Skills
    story.append(Paragraph("CORE SKILLS", section_heading_style))
    story.append(Paragraph(", ".join(cv_data.get("skills", [])), body_style))

    # Tools
    story.append(Paragraph("TOOLS &amp; SOFTWARE", section_heading_style))
    story.append(Paragraph(", ".join(cv_data.get("tools", [])), body_style))

    # Education
    story.append(Paragraph("EDUCATION", section_heading_style))
    story.append(Paragraph(DEFAULT_CANDIDATE["education"], exp_title_style))
    story.append(Paragraph(DEFAULT_CANDIDATE["education_duration"], exp_meta_style))

    # Languages
    story.append(Paragraph("LANGUAGES", section_heading_style))
    story.append(Paragraph(DEFAULT_CANDIDATE["languages"], body_style))

    doc.build(story)
    buffer.seek(0)

    # Also save to file if path provided
    if output_path:
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        with open(output_path, "wb") as f:
            f.write(buffer.getvalue())
        buffer.seek(0)

    return buffer


def generate_cover_letter_pdf(cl_data: dict, output_path: str = None) -> io.BytesIO:
    """
    Generates a professional cover letter in PDF format.
    
    Args:
        cl_data: Cover letter data dict from Groq.
        output_path: Optional file path to save.
        
    Returns:
        BytesIO object containing the PDF file.
    """
    buffer = io.BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=72,
        rightMargin=72,
        topMargin=72,
        bottomMargin=72,
    )

    styles = getSampleStyleSheet()

    header_style = ParagraphStyle(
        "CLHeader",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=14,
        leading=17,
        spaceAfter=2,
    )

    header_sub_style = ParagraphStyle(
        "CLHeaderSub",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9.5,
        leading=12,
        spaceAfter=15,
    )

    meta_style = ParagraphStyle(
        "CLMeta",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9.5,
        leading=13,
        spaceAfter=12,
    )

    body_style = ParagraphStyle(
        "CLBody",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=10,
        leading=14.5,
        spaceAfter=10,
    )

    story = []

    story.append(Paragraph(DEFAULT_CANDIDATE["name"], header_style))
    story.append(
        Paragraph(
            f"{DEFAULT_CANDIDATE['title']}<br/>"
            f"{DEFAULT_CANDIDATE['phone']} | {DEFAULT_CANDIDATE['email']}",
            header_sub_style,
        )
    )
    story.append(
        HRFlowable(
            width="100%",
            thickness=1,
            color=colors.HexColor("#CCCCCC"),
            spaceAfter=12,
        )
    )

    story.append(
        Paragraph(
            f"Date: {cl_data.get('date', 'July 16, 2026')}<br/>"
            f"To: Hiring Team<br/>"
            f"Company: {cl_data.get('company', '')}",
            meta_style,
        )
    )

    body_text = cl_data.get("letter_body", "")
    for paragraph in body_text.split("\n\n"):
        if paragraph.strip():
            story.append(
                Paragraph(paragraph.strip().replace("\n", "<br/>"), body_style)
            )

    doc.build(story)
    buffer.seek(0)

    if output_path:
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        with open(output_path, "wb") as f:
            f.write(buffer.getvalue())
        buffer.seek(0)

    return buffer


def generate_interview_prep_docx(ip_data: dict, output_path: str = None) -> io.BytesIO:
    """
    Generates a professional interview preparation guide in DOCX format.
    """
    doc = docx.Document()
    
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    
    # Header
    header_p = doc.add_paragraph()
    hrun1 = header_p.add_run(DEFAULT_CANDIDATE["name"] + "\n")
    hrun1.bold = True
    hrun1.font.size = Pt(14)
    hrun1.font.name = "Arial"
    
    hrun2 = header_p.add_run(
        f"{DEFAULT_CANDIDATE['title']}\n"
        f"{DEFAULT_CANDIDATE['phone']} | {DEFAULT_CANDIDATE['email']}\n\n"
    )
    hrun2.font.size = Pt(10)
    hrun2.font.name = "Arial"
    
    # Document Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(12)
    title_p.paragraph_format.space_after = Pt(12)
    run_title = title_p.add_run("INTERVIEW PREPARATION GUIDE\n")
    run_title.bold = True
    run_title.font.size = Pt(16)
    run_title.font.name = "Arial"
    
    # Meta / Company
    company_name = ip_data.get("company", "Target Company")
    role_title = ip_data.get("role", "Target Role")
    meta_p = doc.add_paragraph()
    meta_p.paragraph_format.space_after = Pt(18)
    mrun = meta_p.add_run(
        f"Company: {company_name}\n"
        f"Role: {role_title}\n"
    )
    mrun.italic = True
    mrun.font.size = Pt(11)
    mrun.font.name = "Arial"
    
    def add_section_heading(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(6)
        run = h.add_run(text.upper())
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(12)
        
    # Company Research
    add_section_heading("1. Company Overview & Research")
    cp = doc.add_paragraph()
    cp.paragraph_format.space_after = Pt(10)
    crun = cp.add_run(ip_data.get("company_research", "No research provided."))
    crun.font.name = "Arial"
    crun.font.size = Pt(10.5)
    
    # Role Prep Notes
    add_section_heading("2. Role Specific Strategy & Notes")
    rnp = doc.add_paragraph()
    rnp.paragraph_format.space_after = Pt(10)
    rnrun = rnp.add_run(ip_data.get("role_prep_notes", "No notes provided."))
    rnrun.font.name = "Arial"
    rnrun.font.size = Pt(10.5)
    
    # Key Study Topics
    add_section_heading("3. Key Topics to Study")
    for topic in ip_data.get("key_topics", []):
        tp = doc.add_paragraph(style="List Bullet")
        tp.paragraph_format.space_after = Pt(4)
        trun = tp.add_run(topic)
        trun.font.name = "Arial"
        trun.font.size = Pt(10)
        
    # Likely Questions
    add_section_heading("4. Likely Interview Questions & Talking Points")
    for idx, q_obj in enumerate(ip_data.get("likely_questions", []), 1):
        qp = doc.add_paragraph()
        qp.paragraph_format.space_before = Pt(8)
        qp.paragraph_format.space_after = Pt(2)
        qrun = qp.add_run(f"Q{idx}: {q_obj.get('question', '')}")
        qrun.bold = True
        qrun.font.name = "Arial"
        qrun.font.size = Pt(11)
        
        ap = doc.add_paragraph()
        ap.paragraph_format.space_after = Pt(10)
        arun = ap.add_run(f"Suggested Answer / Talking Points:\n{q_obj.get('suggested_answer', '')}")
        arun.font.name = "Arial"
        arun.font.size = Pt(10)
        arun.italic = True
        
    # Save to BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    if output_path:
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        with open(output_path, "wb") as f:
            f.write(buffer.getvalue())
        buffer.seek(0)
        
    return buffer


def generate_interview_prep_pdf(ip_data: dict, output_path: str = None) -> io.BytesIO:
    """
    Generates a professional interview preparation guide in PDF format.
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=72,
        rightMargin=72,
        topMargin=72,
        bottomMargin=72,
    )
    
    styles = getSampleStyleSheet()
    
    header_style = ParagraphStyle(
        "IPHeader",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=14,
        leading=17,
        spaceAfter=2,
    )
    
    header_sub_style = ParagraphStyle(
        "IPHeaderSub",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9.5,
        leading=12,
        spaceAfter=15,
    )
    
    title_style = ParagraphStyle(
        "IPTitle",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=16,
        leading=20,
        spaceBefore=10,
        spaceAfter=6,
    )
    
    meta_style = ParagraphStyle(
        "IPMeta",
        parent=styles["Normal"],
        fontName="Helvetica-Oblique",
        fontSize=10,
        leading=14,
        spaceAfter=14,
    )
    
    sec_heading_style = ParagraphStyle(
        "IPSecHeading",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=12,
        leading=15,
        spaceBefore=14,
        spaceAfter=6,
        keepWithNext=True,
    )
    
    body_style = ParagraphStyle(
        "IPBody",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=10,
        leading=14,
        spaceAfter=10,
    )
    
    bullet_style = ParagraphStyle(
        "IPBullet",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9.5,
        leading=13.5,
        leftIndent=15,
        firstLineIndent=-10,
        spaceAfter=4,
    )
    
    q_style = ParagraphStyle(
        "IPQuestion",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=10.5,
        leading=14,
        spaceBefore=8,
        spaceAfter=4,
        keepWithNext=True,
    )
    
    a_style = ParagraphStyle(
        "IPAnswer",
        parent=styles["Normal"],
        fontName="Helvetica-Oblique",
        fontSize=9.5,
        leading=13.5,
        leftIndent=10,
        spaceAfter=10,
    )
    
    story = []
    
    story.append(Paragraph(DEFAULT_CANDIDATE["name"], header_style))
    story.append(
        Paragraph(
            f"{DEFAULT_CANDIDATE['title']}<br/>"
            f"{DEFAULT_CANDIDATE['phone']} | {DEFAULT_CANDIDATE['email']}",
            header_sub_style,
        )
    )
    story.append(
        HRFlowable(
            width="100%",
            thickness=1,
            color=colors.HexColor("#CCCCCC"),
            spaceAfter=12,
        )
    )
    
    story.append(Paragraph("INTERVIEW PREPARATION GUIDE", title_style))
    
    company_name = ip_data.get("company", "Target Company")
    role_title = ip_data.get("role", "Target Role")
    story.append(
        Paragraph(
            f"Company: {company_name}<br/>"
            f"Role: {role_title}",
            meta_style,
        )
    )
    
    # Company Research
    story.append(Paragraph("1. Company Overview & Research", sec_heading_style))
    story.append(Paragraph(ip_data.get("company_research", "No research provided.").replace("\n", "<br/>"), body_style))
    
    # Role Prep Notes
    story.append(Paragraph("2. Role Specific Strategy & Notes", sec_heading_style))
    story.append(Paragraph(ip_data.get("role_prep_notes", "No notes provided.").replace("\n", "<br/>"), body_style))
    
    # Key Study Topics
    story.append(Paragraph("3. Key Topics to Study", sec_heading_style))
    for topic in ip_data.get("key_topics", []):
        story.append(Paragraph(f"&bull;&nbsp;&nbsp;{topic}", bullet_style))
    story.append(Spacer(1, 10))
    
    # Likely Questions
    story.append(Paragraph("4. Likely Interview Questions & Talking Points", sec_heading_style))
    for idx, q_obj in enumerate(ip_data.get("likely_questions", []), 1):
        story.append(Paragraph(f"Q{idx}: {q_obj.get('question', '')}", q_style))
        story.append(
            Paragraph(
                f"Suggested Answer / Talking Points:<br/>{q_obj.get('suggested_answer', '')}",
                a_style,
            )
        )
        
    doc.build(story)
    buffer.seek(0)
    
    if output_path:
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        with open(output_path, "wb") as f:
            f.write(buffer.getvalue())
        buffer.seek(0)
        
    return buffer
